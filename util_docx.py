"""
Generador de documentos DOCX para reportes del SBE.
Utiliza python-docx para crear documentos profesionales.
"""
import os
from datetime import datetime
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

def _get_downloads_folder():
    """Retorna la ruta de la carpeta Descargas del usuario en Windows."""
    return os.path.join(os.path.expanduser('~'), 'Downloads')

def _add_official_header(doc, title_sub):
    """Añade un encabezado oficial con el logo a la izquierda y texto a la derecha."""
    from database.crud_usuario import listar_instituciones
    inst = None
    try:
        insts = listar_instituciones()
        if insts: inst = insts[0]
    except Exception:
        pass
    
    nombre_inst = inst.get("nombre_institucion", "Institución Educativa") if inst else "Institución Educativa"
    logo_ruta = inst.get("logo_ruta") if inst else None

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(5.0)
    
    cell_logo = table.cell(0, 0)
    cell_text = table.cell(0, 1)
    
    if logo_ruta and os.path.exists(logo_ruta):
        p_logo = cell_logo.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_logo = p_logo.add_run()
        try:
            run_logo.add_picture(logo_ruta, width=Inches(1.2))
        except Exception:
            pass
            
    p_text = cell_text.paragraphs[0]
    p_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run_inst = p_text.add_run(f"República Bolivariana de Venezuela\nMinisterio del Poder Popular para la Educación\n{nombre_inst}\n")
    run_inst.bold = True
    run_inst.font.size = Pt(11)
    
    run_sys = p_text.add_run("SISTEMA DE BRIGADAS ESCOLARES (SBE)\n")
    run_sys.bold = True
    run_sys.font.size = Pt(12)
    run_sys.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
    
    run_sub = p_text.add_run(title_sub)
    run_sub.bold = True
    run_sub.font.size = Pt(14)
    
    doc.add_paragraph("_" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

# =====================================================================
# REPORTES DE INCIDENTES (DOCX)
# =====================================================================

def generar_reporte_docx(reporte: dict, save_path: str = None) -> str | None:
    """
    Recibe un diccionario con los datos del reporte de incidente y genera un archivo DOCX.
    Devuelve la ruta absoluta del archivo generado o None si falla.
    """
    if Document is None:
        print("Error: python-docx no está instalado.")
        return None

    try:
        doc = Document()

        # Configuración de estilos globales
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # -- Encabezado Principal --
        _add_official_header(doc, "REPORTE OFICIAL DE INCIDENTE")

        # -- Información Principal --
        p_info = doc.add_paragraph()
        p_info.add_run("ID Reporte: ").bold = True
        p_info.add_run(f"#{reporte.get('id', 'N/A')}\n")
        
        p_info.add_run("Fecha del Reporte: ").bold = True
        fecha = reporte.get('fecha')
        fecha_str = fecha.strftime("%d/%m/%Y %H:%M") if hasattr(fecha, 'strftime') else str(fecha)
        p_info.add_run(f"{fecha_str}\n")
        
        p_info.add_run("Brigada Asignada: ").bold = True
        p_info.add_run(f"{reporte.get('brigada', 'N/A')}\n")

        p_info.add_run("Estado Actual: ").bold = True
        p_estado = p_info.add_run(f"{reporte.get('estado', 'N/A')}")
        if reporte.get('estado') == 'Resuelto':
            p_estado.font.color.rgb = RGBColor(0x10, 0xb9, 0x81)
        elif reporte.get('estado') == 'En Proceso':
            p_estado.font.color.rgb = RGBColor(0xf5, 0x9e, 0x0b)
        else:
            p_estado.font.color.rgb = RGBColor(0xef, 0x44, 0x44)

        doc.add_paragraph()

        # -- Detalles del Incidente --
        doc.add_heading('Detalles del Incidente', level=2)
        
        p_det = doc.add_paragraph()
        p_det.add_run("Título: ").bold = True
        p_det.add_run(f"{reporte.get('titulo', 'N/A')}\n")
        
        p_det.add_run("Severidad: ").bold = True
        p_prio = p_det.add_run(f"{reporte.get('prioridad', 'N/A')}\n")
        if "Alta" in reporte.get('prioridad', ''):
            p_prio.font.color.rgb = RGBColor(0xef, 0x44, 0x44)
            
        p_det.add_run("Ubicación: ").bold = True
        p_det.add_run(f"{reporte.get('ubicacion', 'N/A')}\n")

        # -- Descripción --
        doc.add_heading('Descripción de los Hechos', level=2)
        p_desc = doc.add_paragraph(reporte.get('descripcion', 'Sin descripción.'))
        p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph()
        doc.add_paragraph("_" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Firma
        doc.add_paragraph()
        doc.add_paragraph()
        p_firma = doc.add_paragraph("_____________________________\nFirma del Responsable / Coordinador")
        p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Pie de página
        footer = doc.add_paragraph(f"\nDocumento generado automáticamente el {datetime.now().strftime('%d/%m/%Y a las %H:%M')}")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

        # -- Guardar Archivo --
        if not save_path:
            filename = f"Reporte_Incidente_{reporte.get('id', 'X')}_{datetime.now().strftime('%Y%m%d')}.docx"
            save_path = os.path.join(_get_downloads_folder(), filename)
            
            counter = 1
            while os.path.exists(save_path):
                filename = f"Reporte_Incidente_{reporte.get('id', 'X')}_{datetime.now().strftime('%Y%m%d')}_{counter}.docx"
                save_path = os.path.join(_get_downloads_folder(), filename)
                counter += 1

        doc.save(save_path)
        return save_path

    except Exception as e:
        print(f"Error generando DOCX: {e}")
        return None

# =====================================================================
# REPORTES DE ACTIVIDADES (DOCX)
# =====================================================================

def generar_reporte_actividad_docx(reporte: dict, save_path: str = None) -> str | None:
    """Genera un reporte DOCX para las Actividades de brigadas escolares."""
    if Document is None: return None
    from database.crud_reporte import obtener_media_reporte
    try:
        doc = Document()
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)

        # -- Encabezado Principal --
        _add_official_header(doc, "INFORME FINAL DE ACTIVIDAD")

        # Datos Generales
        fecha_dt = reporte.get('fecha_reporte')
        fecha_str = fecha_dt.strftime("%d/%m/%Y %H:%M") if hasattr(fecha_dt, 'strftime') else str(fecha_dt)
        act_dt = reporte.get('actividad_fecha')
        act_str = act_dt.strftime("%d/%m/%Y") if hasattr(act_dt, 'strftime') else str(act_dt)

        p_info = doc.add_paragraph()
        p_info.add_run("ID Reporte: ").bold = True
        p_info.add_run(f"ACT-{reporte.get('id', 'N/A')}\n")
        p_info.add_run("Fecha de Emisión: ").bold = True
        p_info.add_run(f"{fecha_str}\n")
        p_info.add_run("Reportado Por: ").bold = True
        p_info.add_run(f"{reporte.get('usuario_nombre', 'N/A')}\n")
        doc.add_paragraph()

        # Detalles de la Actividad
        doc.add_heading('Detalles de la Actividad', level=2)
        p_det = doc.add_paragraph()
        p_det.add_run("Actividad: ").bold = True
        p_det.add_run(f"{reporte.get('actividad_titulo', 'N/A')}\n")
        p_det.add_run("Fecha de Ejecución: ").bold = True
        p_det.add_run(f"{act_str}\n")
        p_det.add_run("Resultado General: ").bold = True
        p_det.add_run(f"{reporte.get('resultado', 'N/A')}\n")

        # Participantes
        participantes = reporte.get('participantes', '')
        if participantes:
            p_det.add_run("Participantes: ").bold = True
            p_det.add_run(f"{participantes}\n")

        # Observaciones
        doc.add_heading('Observaciones', level=2)
        p_res = doc.add_paragraph(reporte.get('resumen', 'Sin observaciones.'))
        p_res.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Multimedia
        archivos = obtener_media_reporte("actividad", reporte.get("id"))
        imagenes = [m for m in archivos if m["tipo"] == "imagen"]
        videos = [m for m in archivos if m["tipo"] == "video"]
        
        if imagenes or videos:
            doc.add_heading('Evidencia Multimedia', level=2)
            if videos:
                p_vid = doc.add_paragraph("Nota: Este reporte incluye videos que no pueden ser incrustados en este documento.")
                p_vid.runs[0].italic = True
                
            for img in imagenes:
                try:
                    doc.add_picture(img["ruta"], width=Inches(5))
                    doc.add_paragraph()  # Espaciado
                except Exception as e:
                    doc.add_paragraph(f"[Error cargando imagen: {img['ruta']}]")

        doc.add_paragraph()
        doc.add_paragraph("_" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_firma = doc.add_paragraph("_____________________________\nFirma del Responsable / Coordinador")
        p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if not save_path:
            filename = f"Reporte_Actividad_ACT-{reporte.get('id', 'X')}_{datetime.now().strftime('%Y%m%d')}.docx"
            save_path = os.path.join(_get_downloads_folder(), filename)
            counter = 1
            while os.path.exists(save_path):
                filename = f"Reporte_Actividad_ACT-{reporte.get('id', 'X')}_{datetime.now().strftime('%Y%m%d')}_{counter}.docx"
                save_path = os.path.join(_get_downloads_folder(), filename)
                counter += 1

        doc.save(save_path)
        return save_path
    except Exception as e:
        print(f"Error generando DOCX de Actividad: {e}")
        return None

# =====================================================================
# REPORTES DE IMPACTO (DOCX)
# =====================================================================

def generar_reporte_impacto_docx(reporte: dict, save_path: str = None) -> str | None:
    """Genera un reporte DOCX para Evaluación de Impacto."""
    if Document is None: return None
    from database.crud_reporte import obtener_media_reporte
    try:
        doc = Document()
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)

        # -- Encabezado Principal --
        _add_official_header(doc, "EVALUACIÓN DE IMPACTO E INDICADORES")

        # Datos Generales
        fecha_dt = reporte.get('fecha_generacion')
        fecha_str = fecha_dt.strftime("%d/%m/%Y %H:%M") if hasattr(fecha_dt, 'strftime') else str(fecha_dt)

        p_info = doc.add_paragraph()
        p_info.add_run("ID Evaluación: ").bold = True
        p_info.add_run(f"IMP-{reporte.get('id', 'N/A')}\n")
        p_info.add_run("Fecha de Evaluación: ").bold = True
        p_info.add_run(f"{fecha_str}\n")
        p_info.add_run("Evaluador: ").bold = True
        p_info.add_run(f"{reporte.get('usuario_nombre', 'N/A')}\n")
        doc.add_paragraph()

        # Datos del Impacto
        doc.add_heading('Datos del Impacto', level=2)
        p_det = doc.add_paragraph()

        if reporte.get('brigada'):
            p_det.add_run("Brigada: ").bold = True
            p_det.add_run(f"{reporte['brigada']}\n")

        if reporte.get('area_evaluada'):
            p_det.add_run("Área Evaluada: ").bold = True
            p_det.add_run(f"{reporte['area_evaluada']}\n")

        if reporte.get('indicador'):
            p_det.add_run("Indicador: ").bold = True
            ind_txt = reporte['indicador']
            if reporte.get('valor'):
                ind_txt += f" — {reporte['valor']}"
            if reporte.get('unidad'):
                ind_txt += f" {reporte['unidad']}"
            p_det.add_run(f"{ind_txt}\n")

        if reporte.get('actividad_titulo'):
            p_det.add_run("Actividad Asociada: ").bold = True
            p_det.add_run(f"{reporte['actividad_titulo']}\n")

        # Descripción del Impacto
        contenido = reporte.get('contenido', '')
        if contenido:
            doc.add_heading('Descripción del Impacto', level=2)
        p_desc = doc.add_paragraph(reporte.get('contenido', 'Sin descripción detallada.'))
        p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Multimedia
        archivos = obtener_media_reporte("impacto", reporte.get("id"))
        imagenes = [m for m in archivos if m["tipo"] == "imagen"]
        videos = [m for m in archivos if m["tipo"] == "video"]
        
        if imagenes or videos:
            doc.add_heading('Evidencia Multimedia', level=2)
            if videos:
                p_vid = doc.add_paragraph("Nota: Este reporte incluye videos que no pueden ser incrustados en este documento.")
                p_vid.runs[0].italic = True
                
            for img in imagenes:
                try:
                    doc.add_picture(img["ruta"], width=Inches(5))
                    doc.add_paragraph()  # Espaciado
                except Exception as e:
                    doc.add_paragraph(f"[Error cargando imagen: {img['ruta']}]")

        doc.add_paragraph()
        doc.add_paragraph("_" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_firma = doc.add_paragraph("_____________________________\nFirma del Evaluador / Director")
        p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if not save_path:
            filename = f"Reporte_Impacto_IMP-{reporte.get('id', 'X')}_{datetime.now().strftime('%Y%m%d')}.docx"
            save_path = os.path.join(_get_downloads_folder(), filename)
            counter = 1
            while os.path.exists(save_path):
                filename = f"Reporte_Impacto_IMP-{reporte.get('id', 'X')}_{datetime.now().strftime('%Y%m%d')}_{counter}.docx"
                save_path = os.path.join(_get_downloads_folder(), filename)
                counter += 1

        doc.save(save_path)
        return save_path
    except Exception as e:
        print(f"Error generando DOCX de Impacto: {e}")
        return None
